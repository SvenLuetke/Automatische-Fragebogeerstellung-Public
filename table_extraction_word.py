from docx import Document
from docx.enum.text import WD_UNDERLINE

def extract_all_tables_data(file_path, skip_first_table=True):
    
    all_tables_data = []

    try:
        document = Document(file_path)
        
        if skip_first_table:
            tables_to_process = document.tables[1:]
        else:
            tables_to_process = document.tables
            
        for table in tables_to_process:
            table_data = []
            
            for row in table.rows:
                row_data = []
                
                for cell in row.cells:
                    cell_html_content = ""
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            text = run.text
                            
                            is_underlined = run.underline not in (False, None, WD_UNDERLINE.NONE)
                            
                            if text:
                                if is_underlined:
                                    cell_html_content += f"<u>{text}</u>"
                                else:
                                    cell_html_content += text


                    
                    row_data.append(cell_html_content.strip())
                    
                table_data.append(row_data)
                
            all_tables_data.append(table_data)
            
    except Exception as e:
        print(f"An error occurred while processing the document: {e}")
        return []

    return all_tables_data

